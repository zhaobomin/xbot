from __future__ import annotations

from pathlib import Path

from ..models import DocumentSection, ParsedDocument, ScannedFile


def parse_docx(path: Path, scanned_file: ScannedFile) -> ParsedDocument:
    import docx

    doc = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    sections = [DocumentSection(text="\n".join(paragraphs))] if paragraphs else []
    return ParsedDocument(
        source_path=scanned_file.path,
        doc_type=scanned_file.doc_type,
        title=paragraphs[0] if paragraphs else path.stem,
        sections=sections,
        modified_time=scanned_file.modified_time,
        content_hash=scanned_file.content_hash,
    )
